from docx import Document
from docx.shared import Pt,RGBColor
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from app.docxutils.TextCorrection import textcorrectApi
import re
from operator import itemgetter,attrgetter
def text_process(path):
    document = Document(path)
    paragraphs = document.paragraphs
    document.add_page_break()
    Head = document.add_paragraph()
    h1=Head.add_run('===---***错误报告***---===')
    h1.blod=True
    h1.font.size=Pt(25)
    h1.font.color.rgb = RGBColor(222,203,0)
    Head.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    error_word=0
    paragraph_num=0
    for paragraph in paragraphs:
        paragraph_num+=1
        word_num=0
        sentences = re.split(r"([。!！?？\n])", paragraph.text)
        sentences.append("")
        sentences = ["".join(i) for i in zip(sentences[0::2], sentences[1::2])]
        for word in sentences:
            if(len(word.strip())>=1):#移除首字符后判断该字符串是否是空的
                word_num+=1
                correct_dict=eval(textcorrectApi(word))
                print(type(correct_dict))
                errormessage=[]
                for black_list in correct_dict['black_list']:
                    errormessage.append(black_list)
                for punc_correct in correct_dict['punc']:
                    errormessage.append(punc_correct)
                for leader_correct in correct_dict['leader']:
                    errormessage.append(leader_correct)
                for org_correct in correct_dict['org']:
                    errormessage.append(org_correct)
                for pol_correct in correct_dict['pol']:
                    errormessage.append(pol_correct)
                for grammar_pc_correct in correct_dict['grammar_pc']:
                    errormessage.append(grammar_pc_correct)
                for order_correct in correct_dict['order']:
                    errormessage.append(order_correct)
                for idm_correct in correct_dict['idm']:
                    errormessage.append(idm_correct)
                for word_correct in correct_dict['word']:
                    errormessage.append(word_correct)
                for char_correct in correct_dict['char']:
                    errormessage.append(char_correct)
                for redund_correct in correct_dict['redund']:
                    errormessage.append(redund_correct)
                for miss_correct in correct_dict['miss']:
                    errormessage.append(miss_correct)
                for dapei_correct in correct_dict['dapei']:
                    errormessage.append(dapei_correct)
                for number_correct in correct_dict['number']:
                    errormessage.append(number_correct)
                for addr_correct in correct_dict['addr']:
                    errormessage.append(addr_correct)
                for name_correct in correct_dict['name']:
                    errormessage.append(name_correct)
                if len(errormessage)>=1:
                    error_word+=1
                    partWord = document.add_paragraph()
                    pw1=partWord.add_run('==--**第'+str(error_word)+'处错误句子**--==')
                    pw1.blod=True
                    pw1.font.size=Pt(20)
                    pw1.font.color.rgb = RGBColor(222,203,0)
                    paragraph=document.add_paragraph()
                    p1=paragraph.add_run("原错误句子：")
                    p1.font.size=Pt(16)
                    p1.blod=True
                    p1.font.color.rgb=RGBColor(161,90,56)
                    p2=paragraph.add_run(word)
                    p2.font.size=Pt(16)
                    corrdinate = document.add_paragraph()
                    c1=corrdinate.add_run("句子坐标：")
                    c1.blod=True
                    c1.font.size=Pt(16)
                    c1.font.color.rgb = RGBColor(43,87,154)
                    c2=corrdinate.add_run("第")
                    c2.font.size=Pt(16)
                    c3=corrdinate.add_run(str(paragraph_num))
                    c3.font.size=Pt(16)
                    c3.font.color.rgb = RGBColor(43,87,184)
                    c4=corrdinate.add_run("段,第")
                    c4.font.size=Pt(16)
                    c5=corrdinate.add_run(str(word_num))
                    c5.font.size=Pt(16)
                    c5.font.color.rgb = RGBColor(43,87,184)
                    c6=corrdinate.add_run("句")
                    c6.font.size=Pt(16)
                    i=0
                    for errormessage_part in errormessage:
                        i+=1
                        paragraph1=document.add_paragraph()
                        p11=paragraph1.add_run("第"+str(i)+"处出错地方：")
                        p11.font.size=Pt(12)
                        p11.blod=True
                        paragraph1.add_run(word[0:errormessage_part[0]]).font.size=Pt(12)
                        pc=paragraph1.add_run(word[errormessage_part[0]:errormessage_part[0]+len(errormessage_part[1])])
                        pc.font.color.rgb=RGBColor(234,49,49)
                        pc.font.size=Pt(15)
                        paragraph1.add_run(word[errormessage_part[0]+len(errormessage_part[1]):]).font.size=Pt(12)
                        paragraph2=document.add_paragraph()
                        p21=paragraph2.add_run('该句第'+str(errormessage_part[0])+'个字符处---建议将"'+errormessage_part[1]+'"改为"'+errormessage_part[2]+'"')
                        p21.font.size=Pt(12)
                        p21.blod=True
                    paragraph3=document.add_paragraph()
                    p31=paragraph3.add_run("该句修改后为：")
                    p31.font.size=Pt(16)
                    p31.blod=True
                    for i in range(1,len(errormessage)):
                        for j in range(0,len(errormessage)-i):
                            if errormessage[j][0]>errormessage[j+1][0]:
                                errormessage[j],errormessage[j+1]=errormessage[j+1],errormessage[j]
                    for index in range(len(errormessage)-1,-1,-1):
                        error=errormessage[index]
                        word=word[0:error[0]]+error[2]+word[error[0]+len(error[1]):]
                    p32=paragraph3.add_run(word)
                    p32.font.size=Pt(16)
                    p32.font.underline = True
    document.save(path)
    return 100

def get_file_byte(filename):
    with open(filename, "rb") as f:
        while True:
            content = f.read(1024)
            if content:
                yield content
            else:
                break
