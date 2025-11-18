import difflib
import json
from docx import Document
def do_com(path1,path2):
    text1=""
    text2=""
    document1=Document(path1)
    document2=Document(path2)
    paragraphs1=document1.paragraphs
    paragraphs2=document2.paragraphs
    for pa1 in paragraphs1:
        text1=text1+pa1.text
    for pa2 in paragraphs2:
        text2=text2+pa2.text
    sequenceMatcher=difflib.SequenceMatcher()
    sequenceMatcher.set_seqs(text1,text2)
    similar_num=sequenceMatcher.ratio()
    print(similar_num)
    # text1=text1.splitlines(keepends=True)
    # text2=text2.splitlines(keepends=True)
    # d = difflib.Differ()
    # fruitWord=list(d.compare(text1,text2))
    # for f in fruitWord:
    #     print(f)
    # fruit=[similar_num]
    # count=0
    # for f in fruitWord:
    #     if f[0]=='+' or f[0]=='-':
    #         word=f[1:]
    #         allcolor=f[0]
    #         addw=[]
    #         delw=[]
    #         dif=[]
    #         if count+1<len(fruitWord):
    #             if fruitWord[count+1][0]=='?':
    #                 temp=-1
    #                 for c in fruitWord[count+1]:
    #                     if c=='+':
    #                         addw.append(temp)
    #                     if c=='-':
    #                         delw.append(temp)
    #                     if c=='^':
    #                         dif.append(temp)
    #                     temp+=1
    #         difword={"word":word,"allcolor":allcolor,"addw":addw,"delw":delw,"dif":dif}
    #         fruit.append(difword)
    #     count+=1
    # jsonstr=json.dumps(fruit)
    # print(jsonstr)
    return similar_num

def gettext(path1):
    text1=""
    document1=Document(path1)
    paragraphs1=document1.paragraphs
    for pa1 in paragraphs1:
        text1=text1+pa1.text+'\n'
    # print(text1)
    return text1;
if __name__=="__main__":
    # text1="你好，我叫汤城冠"
    # text2="你好，我是汤城冠"
    text1=r"E://Project//docx//b305086d-2430-11ef-af12-cc6b1e8cd122.docx"
    text2=r"E://Project//docx//0225e56a-2430-11ef-beb2-cc6b1e8cd122.docx"
    do_com(text1,text2)