import os

from docx import Document
import pandas as pd
import zipfile


def do_bat(path1,path2,path3,zipName,control):
    os.mkdir(path3)
    head_list,target_dic=scan_xlsl(path2)
    print(head_list,target_dic)
    count= 0
    for i in range(0,len(target_dic)):
        document=Document(path1)
        for j in range(0,len(head_list)):
            for paragraph in document.paragraphs:
                for run in paragraph.runs:
                    if head_list[j] in run.text:
                        run.text=run.text.replace(str(head_list[j]),str(target_dic[i][head_list[j]]))
        if control=='number' and control not in head_list:
            print("123")
            document.save(path3+str(count)+'.docx')
        else:
            print("456")
            print(str(control))
            print(str(target_dic[i][control])+"1")
            document.save(path3+str(target_dic[i][control])+'.docx')
        count+=1
    zip = zipfile.ZipFile("C:\\tcg\\MySoftWare\\textprocess\\Project\\zip\\"+zipName+".zip","w",zipfile.ZIP_DEFLATED)
    for path,dirname,filenames in os.walk(path3):
        fpath=path.replace(path3,'')
        for filename in filenames:
            zip.write(os.path.join(path,filename) , os.path.join(fpath,filename))
    zip.close()
    return 100

def scan_xlsl(path):
    ex_data = pd.read_excel(path) #默认读取第一个sheet的内容
    head_list = list(ex_data.columns) #拿到表头: [A, B, C, D]
    list_dic = []
    for i in ex_data.values: # i 为每一行的value的列表：[a2, b2, c2, d2]
        a_line = dict(zip(head_list, i)) # a_line: {"A": "a2", "B": "b2", "C": "c2", "D": "d2"}
        list_dic.append(a_line)
    return head_list,list_dic
# def file_name_bat(path,path2,path3,ZipName,control,copy_num):
#     import shutil
#     print("file_bat")
#     os.mkdir(path3)
#     filetype=path.split('.')[-1]
#     if copy_num!='0':
#         for i in range(int(copy_num)):
#             shutil.copy(path,path3+str(i)+'.'+filetype)
#     else:
#         if control=='txt':
#             pass
#         else:
#             excel_data=pd.read_excel(path2,usecols=[0])
#             list_data=[excel_data.columns[0]]
#             for i in excel_data.values:
#                 list_data.append(i[0])
#         for data in list_data:
#             shutil.copy(path,path3+data+'.'+filetype)
#     zip = zipfile.ZipFile("C:\\tcg\\MySoftWare\\textprocess\\Project\\zip\\"+ZipName+".zip","w",zipfile.ZIP_DEFLATED)
#     for path,dirname,filenames in os.walk(path3):
#         fpath=path.replace(path3,'')
#         for filename in filenames:
#             zip.write(os.path.join(path,filename) , os.path.join(fpath,filename))
#     zip.close()
#     return 100
def get_Similarity(path,Address,Fname):
    # 1 创建工作薄
    import openpyxl
    from openpyxl.styles import Alignment
    from openpyxl.styles import Font
    import difflib
    VersionName = []
    SimilarNum = []
    count=0
    smallSimilar=2.0
    smallSimilarVersionName=""
    bigSimilar=-1.0
    bigSimilarVersionName=""
    text1=""
    document1=Document(path)
    paragraphs1=document1.paragraphs
    for pa1 in paragraphs1:
        text1=text1+pa1.text
    for docx in os.scandir(Address):
        if docx.is_file():
            count=count+1
            versionpath=Address+"/"+docx.name
            print(versionpath)
            document2=Document(versionpath)
            paragraphs2=document2.paragraphs
            text2=""
            for pa2 in paragraphs2:
                text2=text2+pa2.text
            sequenceMatcher=difflib.SequenceMatcher()
            sequenceMatcher.set_seqs(text1,text2)
            similar_num=sequenceMatcher.ratio()
            if similar_num<=smallSimilar:
                smallSimilar=similar_num
                smallSimilarVersionName=docx.name
                smallSimilarVersionName=smallSimilarVersionName.split('.')[0]+"."+smallSimilarVersionName.split('.')[1]
            if similar_num>=bigSimilar:
                bigSimilar=similar_num
                bigSimilarVersionName=docx.name
                bigSimilarVersionName=bigSimilarVersionName.split('.')[0]+"."+bigSimilarVersionName.split('.')[1]
            SimilarNum.append(similar_num)
            VersionName.append(docx.name.split('.')[0]+"."+docx.name.split('.')[1])
    workbook = openpyxl.Workbook()
    # 获取活动工作表， 默认就是第一个工作表
    worksheet = workbook.active
    worksheet.title = "相似度统计结果"
    worksheet.column_dimensions['A'].width = 25.0
    worksheet.column_dimensions['B'].width = 25.0
    worksheet.column_dimensions['C'].width = 25.0
    worksheet.column_dimensions['D'].width = 25.0
    worksheet.merge_cells('A1:D2')
    worksheet.cell(1,1).value='该文件一共和'+str(count)+'个文件进行相似度对比'
    worksheet.merge_cells('A3:D4')
    worksheet.cell(3,1).value='其中与该文件相似度最大的是'+bigSimilarVersionName+',相似度为'+str(bigSimilar)
    worksheet.merge_cells('A5:D6')
    worksheet.cell(5,1).value='其中与该文件相似度最小的是'+smallSimilarVersionName+',相似度为'+str(smallSimilar)
    worksheet.merge_cells('A7:D8')
    worksheet.cell(7,1).value='详情请见“相似度统计表详情”表'
    ac_area=worksheet['A1:D8']
    align=Alignment(horizontal='center',vertical='center')
    for i in ac_area:
        for j in i:
            j.alignment = align
    # 2 以下是我们要写入的数据
    worksheet2=workbook.create_sheet(0)
    worksheet2.title = "相似度统计表详情"
    worksheet2.column_dimensions['A'].width = 14.0
    worksheet2.column_dimensions['B'].width = 17.0
    Project = ['版本号','与该论文的相似度']
    # 写入第一行数据，行号和列号都是从 1 开始计数
    for i in range(len(Project)):
        worksheet2.cell(1, i + 1, Project[i])
    # 写入第一列数据， 因为第一行已经有数据了，所以是 i + 2
    for i in range(len(VersionName)):
        worksheet2.cell(i + 2, 1, VersionName[i])
    # 写入第二列数据
    for i in range(len(SimilarNum)):
        worksheet2.cell(i + 2, 2, SimilarNum[i])
    worksheet2['A1'].font = Font(bold=True)
    worksheet2['B1'].font = Font(bold=True)
    workbook.save(filename="C:\\tcg\\MySoftWare\\textprocess\\Project\\excel\\"+Fname+".xlsx")
    return 100
if __name__ == '__main__':
    get_Similarity(r"E:/Project/user/1234567890000/7c30e479-6b54-4f91-bd76-2d3fcbf6b020/1a69dcdb-ac4d-4783-956c-b859f8b4ba1b.docx",r"E:/Project/user/1111111111111","test")
    # scan_xlsl("E:/TestDir/test.xlsx")
