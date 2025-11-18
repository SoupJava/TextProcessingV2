import re

from docx import Document
from docx.shared import RGBColor, Pt


def do_docx(path):
    document = Document(path)
    paragraphs = document.paragraphs
    report = False
    reports = []
    essay = []
    for paragraph in paragraphs:
        if paragraph.text == '===---***错误报告***---===':
            report = True
        sentences = re.split(r"([。!！?？\n])", paragraph.text)
        sentences.append("")
        sentences = ["".join(i) for i in zip(sentences[0::2], sentences[1::2])]
        if report:
            if sentences[0][0:4] == '原错误句':
                correct_word = sentences[0][6:]
                句子去掉空行前长度=len(correct_word)
                correct_word = correct_word.replace('\t', '')
                correct_word = correct_word.replace(' ', '')
                句子原长度=len(correct_word)
                lack=句子去掉空行前长度-句子原长度
            if sentences[0][0:4] == '句子坐标':
                paragraph_num = re.findall(r'第(.*?)段', sentences[0])
                word_num = re.findall(r',第(.*?)句', sentences[0])
            if sentences[0][0:4] == '该句修改':
                new_word = sentences[0][7:]
                new_word = new_word.replace('\t', '')
                new_word = new_word.replace(' ', '')
                report_message = {'paragraph': paragraph_num, 'word': word_num, 'correct': correct_word,
                                  'new': new_word,'lack':lack}
                reports.append(report_message)
        else:
            for word in sentences:
                essay.append(word)
    for report in reports:
        print(report)
    paragraph_num2 = 0
    count = 0
    control=False
    for paragraph in document.paragraphs:
        word_num2 = 0
        paragraph_num2 += 1
        font_num = 0
        sentences = re.split(r"([。!！?？\n])", paragraph.text)
        sentences.append("")
        sentences = ["".join(i) for i in zip(sentences[0::2], sentences[1::2])]
        print(sentences)
        new_paragraph=""
        for word in sentences:
            print(word)
            word_num2 += 1
            if count<len(reports):
                if paragraph_num2 == int(reports[count]['paragraph'][0]):
                    if word_num2 == int(reports[count]['word'][0]):
                        control=True
                        new_text=""
                        for i in range(0,int(reports[count]['lack'])):
                            new_text+=" "
                        new_paragraph+=new_text+reports[count]['new']
                        count += 1
                    else:
                        new_paragraph+=word
                else:
                    new_paragraph+=word
            else:
                new_paragraph+=word
            font_num += len(word)
        if control:
            text=new_paragraph
            old_paragraph=paragraph
            paragraph.text=""
            cp=paragraph.add_run(text)
            cp.font.size = Pt(15)  # 字号大小
            cp.font.bold = False  # 是否加粗
            cp.font.italic = False  # 是否斜体
            cp.font.underline = False  # 是否下划线
            cp.font.shadow = True  # 是否阴影
            cp.font.color.rgb = RGBColor(222,203,0)  # 字体颜色
            copy_paragraph(old_paragraph,paragraph)
        control=False
    delete_control=False
    for paragraph in document.paragraphs:
        if paragraph.text == '===---***错误报告***---===':
            delete_control = True
        if delete_control:
            delete_paragraph(paragraph)
    document.save(path)
    return 100

def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    p._p = p._element = None

def copy_paragraph(old_paragraph, new_paragraph):
    new_paragraph.paragraph_format.line_spacing=old_paragraph.paragraph_format.line_spacing
    new_paragraph.paragraph_format.first_line_indent=old_paragraph.paragraph_format.first_line_indent
    new_paragraph.paragraph_format.space_before=old_paragraph.paragraph_format.space_before
    new_paragraph.paragraph_format.space_after=old_paragraph.paragraph_format.space_after
    new_paragraph.paragraph_format.alignment=old_paragraph.paragraph_format.alignment


# if __name__ == '__main__':
#     # text="1234532423423"
#     # text=text[0:4]+"ASDSDA"+text[6:]
#     # print(text)
#     do_docx("C://Users//13977//Desktop//下载//Correct (17).docx")
